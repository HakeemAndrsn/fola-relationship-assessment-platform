import os
import sys

def main():
    print("🚀 Initializing python-docx generation script...")
    
    # Try importing docx, if not found, we will guide the user
    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml import OxmlElement, parse_xml
        from docx.oxml.ns import nsdecls, qn
    except ImportError:
        print("❌ python-docx is not installed. Please run via: uv run --with python-docx generate_docx.py")
        sys.exit(1)
        
    output_dir = "/Users/fola/.gemini/antigravity-cli/brain/77500196-dec9-416a-978c-68d6cb4ec0c9"
    cover_image_path = os.path.join(output_dir, "somatic_workbook_cover.jpg")
    
    # Check if cover image exists
    has_cover = os.path.exists(cover_image_path)
    if not has_cover:
        print(f"⚠️ Cover image not found at {cover_image_path}. Generation will proceed without image cover.")

    # Helper function to add customized headings
    def add_custom_heading(doc, text, level, space_before=18, space_after=6, color=(36, 22, 19)):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after = Pt(space_after)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.font.name = "Playfair Display" if level <= 2 else "Inter"
        run.font.size = Pt(24 if level == 1 else 18 if level == 2 else 14)
        run.bold = True
        run.font.color.rgb = RGBColor(*color)
        return p

    # Helper function to add styled body paragraphs
    def add_body_paragraph(doc, text, italic=False, bold_prefix="", color=(71, 85, 105), space_after=10):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(space_after)
        p.paragraph_format.line_spacing = 1.15
        
        if bold_prefix:
            run_bold = p.add_run(bold_prefix)
            run_bold.font.name = "Inter"
            run_bold.bold = True
            run_bold.font.color.rgb = RGBColor(36, 22, 19)
            
        run = p.add_run(text)
        run.font.name = "Inter"
        run.font.size = Pt(11)
        run.italic = italic
        run.font.color.rgb = RGBColor(*color)
        return p

    # Helper function to add journal lines
    def add_journal_lines(doc, num_lines=8):
        for _ in range(num_lines):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            run = p.add_run("__________________________________________________________________")
            run.font.name = "Inter"
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(226, 232, 240)
        
        p_space = doc.add_paragraph()
        p_space.paragraph_format.space_after = Pt(12)

    # Helper function to add highlight boxes
    def add_highlight_box(doc, title, text, color=(107, 39, 55)):
        tbl = doc.add_table(rows=1, cols=1)
        tbl.autofit = False
        tbl.columns[0].width = Inches(6.0)
        
        cell = tbl.cell(0, 0)
        # Background color shading
        shading = parse_xml(r'<w:shd {} w:fill="FAF8F5"/>'.format(nsdecls('w')))
        cell._tc.get_or_add_tcPr().append(shading)
        
        # Left border thickness/color
        tcPr = cell._tc.get_or_add_tcPr()
        tcBorders = parse_xml(r'''
            <w:tcBorders {}>
                <w:top w:val="none"/>
                <w:left w:val="single" w:sz="24" w:space="0" w:color="6B2737"/>
                <w:bottom w:val="none"/>
                <w:right w:val="none"/>
            </w:tcBorders>
        '''.format(nsdecls('w')))
        tcPr.append(tcBorders)
        
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(4)
        run_title = p.add_run(f"{title}\n")
        run_title.bold = True
        run_title.font.name = "Inter"
        run_title.font.size = Pt(11)
        run_title.font.color.rgb = RGBColor(*color)
        
        run_text = p.add_run(text)
        run_text.font.name = "Inter"
        run_text.font.size = Pt(10.5)
        run_text.italic = True
        run_text.font.color.rgb = RGBColor(71, 85, 105)
        
        doc.add_paragraph().paragraph_format.space_after = Pt(12)


    # =========================================================================
    # DOCUMENT 1: SOMATIC COMPANION WORKBOOK
    # =========================================================================
    print("Generating Somatic Companion Workbook Word document...")
    doc1 = Document()
    
    # Margins
    for section in doc1.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    # Cover Art Page
    if has_cover:
        # We add the cover art image
        p = doc1.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run()
        run.add_picture(cover_image_path, width=Inches(6.0))
        doc1.add_page_break()
    else:
        # Backup Cover
        p = doc1.add_paragraph()
        p.paragraph_format.space_before = Pt(100)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_tag = p.add_run("THE LADDERWORK SEQUENCE — STEP 1\n\n")
        run_tag.font.name = "Inter"
        run_tag.font.size = Pt(12)
        run_tag.bold = True
        run_tag.font.color.rgb = RGBColor(201, 185, 154)
        
        run_title = p.add_run("The Inner Child Somatic\nCompanion Journal\n\n")
        run_title.font.name = "Playfair Display"
        run_title.font.size = Pt(36)
        run_title.bold = True
        run_title.font.color.rgb = RGBColor(36, 22, 19)
        
        run_sub = p.add_run("A 7-day somatic workbook to map triggers, locate the childhood origin of defense patterns, and integrate a new state of emotional safety.")
        run_sub.font.name = "Inter"
        run_sub.font.size = Pt(14)
        run_sub.font.color.rgb = RGBColor(107, 39, 55)
        
        doc1.add_page_break()

    # Intro Page
    add_custom_heading(doc1, "Understanding Somatic Traces", 1)
    add_body_paragraph(doc1, "The Core Philosophy of the Journal", italic=True, color=(107, 39, 55), space_after=18)
    
    add_body_paragraph(doc1, "Welcome to your Step 1 Somatic Companion. Your conscious mind is an intellectual strategist. It can write lists, set boundaries, and explain why you feel stuck. However, your survival defenses—your sudden urges to withdraw, your fear of visibility, your anxiety—do not live in your logic. They live in your nervous system.")
    add_body_paragraph(doc1, "When you experience a trigger today (whether a relationship conflict or a professional block), your body is executing a defense script that was written years ago to keep you safe in a child's environment. We call this the Somatic Trace.")
    add_body_paragraph(doc1, "This journal is designed to help you locate that trace, map its childhood address, and document the integration process. Use these pages before and after you play your Step 1 Regression Audio. Write without filtering. Let the subconscious speak.")
    
    add_highlight_box(doc1, "Somatic Axiom:", "You cannot think your way out of a physiological defense response. To shift the state, you must locate the somatic address where the script was first compiled.")
    
    p_disc = doc1.add_paragraph()
    p_disc.paragraph_format.space_before = Pt(50)
    run_disc = p_disc.add_run("Disclaimer: FOLA is a complementary somatic practice. This workbook does not diagnose, treat, or manage medical or clinical mental-health conditions. If you are in crisis, please contact SADAG on 0800 567 567 immediately.")
    run_disc.font.name = "Inter"
    run_disc.font.size = Pt(9)
    run_disc.italic = True
    run_disc.font.color.rgb = RGBColor(148, 163, 184)
    
    doc1.add_page_break()

    # Day 1
    add_custom_heading(doc1, "Day 1: Locating the Somatic Trace", 1)
    add_body_paragraph(doc1, "Identify a moment this past week when you felt defensive, small, or irrationally anxious. Map where it registered in your body. Don't write what you thought—write what you physically felt.", color=(71, 85, 105))
    add_body_paragraph(doc1, "Somatic Locations Checked: [  ] Throat (Tightness, vocal constraint)   [  ] Chest (Heavy weight, shallow breathing)   [  ] Stomach (Fluttering, hollow, drop)", italic=True, color=(107, 39, 55))
    
    add_body_paragraph(doc1, "Prompt: Describe the physical sensations and write the earliest memory where you felt this exact somatic response. Who was in the room? What was happening?", bold_prefix="Journal Prompt: ")
    add_journal_lines(doc1, 8)
    doc1.add_page_break()

    # Day 2
    add_custom_heading(doc1, "Day 2: The Protection Script", 1)
    add_body_paragraph(doc1, "As a child, when your environment felt unsafe or overwhelming, you developed a defense mechanism. Which of these four scripts did you build? (Underline yours: The Avoidant One, The Academic Performer, The Rebel Rouser, The Over-Eager Helper).")
    add_body_paragraph(doc1, "Prompt: How did this script keep you safe when you were 8 years old? How is it costing you money, intimacy, or sleep now that you are an adult?", bold_prefix="Journal Prompt: ")
    add_journal_lines(doc1, 8)
    doc1.add_page_break()

    # Day 3
    add_custom_heading(doc1, "Day 3: Hypnotic Preparation", 1)
    add_body_paragraph(doc1, "Complete this page right before you listen to your Step 1 Regression Audio. Rate your current somatic state. Are you hyper-vigilant (tense/activated) or shut down (exhausted/flat)?")
    add_body_paragraph(doc1, "Nervous System Scale: 1 (Flat/Numb) ----- 4 (Grounded/At Ease) ----- 7 (Activated/Tense) ----- 10 (Panicked/Flight)", italic=True, color=(107, 39, 55))
    add_body_paragraph(doc1, "Prompt: Set your intention. Write a single sentence to your subconscious: 'I am ready to show you that we are safe now.' What does safety feel like in your chest today?", bold_prefix="Journal Prompt: ")
    add_journal_lines(doc1, 8)
    doc1.add_page_break()

    # Day 4
    add_custom_heading(doc1, "Day 4: Post-Audio Integration", 1)
    add_body_paragraph(doc1, "Complete this page immediately after the guided audio. Write while your prefrontal cortex is quiet and the release is fresh.")
    add_body_paragraph(doc1, "Prompt: What raw words did your younger self speak during the regression? What is the concrete, structural boundary your adult self must set this week to honor that child?", bold_prefix="Journal Prompt: ")
    add_journal_lines(doc1, 8)
    doc1.add_page_break()

    # Day 5
    add_custom_heading(doc1, "Day 5: Reclaiming Agency", 1)
    add_body_paragraph(doc1, "Scarcity is a survival strategy. It makes you feel like resources, love, and safety are finite. It drives you to say yes to clients, relationships, and situations that drain you.")
    add_body_paragraph(doc1, "Prompt: If you knew, with absolute physical certainty, that you could not be abandoned or go broke, what choice would you make in your life or business today?", bold_prefix="Journal Prompt: ")
    add_journal_lines(doc1, 8)
    doc1.add_page_break()

    # Day 6
    add_custom_heading(doc1, "Day 6: Relational Traces", 1)
    add_body_paragraph(doc1, "Intimacy triggers the inner child because closeness feels threatening to an early protection script. When your partner pulls away or criticizes you, your nervous system responds with the old threat code.")
    add_body_paragraph(doc1, "Prompt: Map your relationship trigger loop. When they do X, I physically feel Y. Write down a rewired somatic response you can execute next time this loop is triggered.", bold_prefix="Journal Prompt: ")
    add_journal_lines(doc1, 8)
    doc1.add_page_break()

    # Day 7
    add_custom_heading(doc1, "Day 7: Anchor Your Shift", 1)
    add_body_paragraph(doc1, "Subconscious transformation is solidified when it is spoken. By expressing your progress aloud, you wire the new state of safety directly into your nervous system. Complete the 7-day sequence by recording a brief, 2-minute video.")
    
    add_highlight_box(doc1, "Your 2-Minute Reflection Outline:", 
                      "1. Before: How did you feel walking into this Step 1 package? (Anxiety, block, or tension)\n"
                      "2. During: What occurred in your body during the guided regression audio?\n"
                      "3. After: How does your body (chest, throat, stomach) feel right now?\n"
                      "4. Next: What are you looking forward to next on your healing journey?")
                      
    add_body_paragraph(doc1, "Scan the Portal QR Code on your dashboard to upload your reflection. Once uploaded, your dashboard will instantly unlock your R200 LoveBetter Assessment Voucher and email your 10-Min Somatic Emergency Reset Audio.", color=(107, 39, 55))
    
    # Save doc 1
    doc1_path = os.path.join(output_dir, "FOLA_Inner_Child_Somatic_Workbook.docx")
    doc1.save(doc1_path)
    print(f"✅ Somatic Workbook DOCX Successfully Saved to: {doc1_path}")


    # =========================================================================
    # DOCUMENT 2: GUIDED SOMATIC REGRESSION HYPNOSIS SCRIPT
    # =========================================================================
    print("Generating Guided Somatic Regression Hypnosis Script Word document...")
    doc2 = Document()
    
    for section in doc2.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    # Cover Art Page
    if has_cover:
        p = doc2.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run()
        run.add_picture(cover_image_path, width=Inches(6.0))
        doc2.add_page_break()
    else:
        p = doc2.add_paragraph()
        p.paragraph_format.space_before = Pt(100)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_tag = p.add_run("THE LADDERWORK SEQUENCE — STEP 1 AUDIO\n\n")
        run_tag.font.name = "Inter"
        run_tag.font.size = Pt(12)
        run_tag.bold = True
        run_tag.font.color.rgb = RGBColor(201, 185, 154)
        
        run_title = p.add_run("Guided Somatic Regression\nHypnosis Script\n\n")
        run_title.font.name = "Playfair Display"
        run_title.font.size = Pt(36)
        run_title.bold = True
        run_title.font.color.rgb = RGBColor(36, 22, 19)
        
        run_sub = p.add_run("The official recording script for Hakeem Lesolang. Step 1 Audio guide to accompany the Somatic Companion Journal.")
        run_sub.font.name = "Inter"
        run_sub.font.size = Pt(14)
        run_sub.font.color.rgb = RGBColor(107, 39, 55)
        
        doc2.add_page_break()

    # About the Recording
    add_custom_heading(doc2, "Practitioner Production Guidelines", 1)
    add_body_paragraph(doc2, "Audio Recording Instructions for Hakeem Lesolang", italic=True, color=(107, 39, 55))
    
    add_body_paragraph(doc2, "Pacing and Tone:", bold_prefix="1. ")
    add_body_paragraph(doc2, "The voice must be grounded, warm, and highly calm. Use a slow, rhythmic cadence. Speak in a baroque tempo—roughly 60 beats per minute—matching speech patterns to natural exhalations. Allow for 2-3 second silences between instructions to let the client sink deeper.", color=(71, 85, 105))
    
    add_body_paragraph(doc2, "Sound Frequencies & Background:", bold_prefix="2. ")
    add_body_paragraph(doc2, "Layer a soft 396Hz frequency underneath the recording. This frequency is traditionally associated with releasing fear, guilt, and ancestral safety locks. Add a low-volume pink noise base to mask room room echo and normalize headphone listening.", color=(71, 85, 105))
    
    add_body_paragraph(doc2, "Target Audience:", bold_prefix="3. ")
    add_body_paragraph(doc2, "High-performing Black professionals, founders, and executives who are intellectually accomplished but somatic vigilance is high. The language must feel premium, non-clinical, and somatic.", color=(71, 85, 105))
    
    doc2.add_page_break()

    # The Script
    add_custom_heading(doc2, "Recording Script: Guided Somatic Regression", 1)
    add_body_paragraph(doc2, "Step 1 Audio Guide — ~20 Minutes", italic=True, color=(107, 39, 55))
    
    add_custom_heading(doc2, "Section 1: Somatic Induction (0:00 - 4:00)", 2)
    add_highlight_box(doc2, "Speaking Notes", "Slow, deep, anchoring tone. Match the exhalations.")
    
    add_body_paragraph(doc2, "Sit comfortably. Close your eyes. Drop your shoulders away from your ears. Let the weight of your body sink fully into the chair beneath you.", italic=True, color=(36, 22, 19))
    add_body_paragraph(doc2, "You have spent the day managing, thinking, executing. For the next twenty minutes, there is nothing you need to fix. There is no one you need to save. Your nervous system is allowed to rest.", italic=True, color=(36, 22, 19))
    add_body_paragraph(doc2, "Take a deep breath in through your nose... hold it... and release it through your mouth with a soft sigh. Let your chest drop. Good.", italic=True, color=(36, 22, 19))
    add_body_paragraph(doc2, "As you listen to the sound of my voice, feel a wave of calm relaxation starting at the top of your head, sliding down behind your eyes, relaxing your jaw, settling into your throat, and softening your chest. Every breath out takes you deeper into a peaceful, focused state of authority.", italic=True, color=(36, 22, 19))
    
    add_custom_heading(doc2, "Section 2: Bypassing the Critical Factor (4:00 - 8:00)", 2)
    add_highlight_box(doc2, "Speaking Notes", "Rhythmic counting. Leave space between numbers.")
    
    add_body_paragraph(doc2, "Your conscious mind is a brilliant protector. But right now, we invite it to step back. We thank it for its service, and we let it rest.", italic=True, color=(36, 22, 19))
    add_body_paragraph(doc2, "Imagine a staircase of five steps leading down into your own safe, warm, subconscious vault.", italic=True, color=(36, 22, 19))
    add_body_paragraph(doc2, "Five... taking you deeper.", italic=True, color=(36, 22, 19))
    add_body_paragraph(doc2, "Four... letting go of external noise.", italic=True, color=(36, 22, 19))
    add_body_paragraph(doc2, "Three... feeling more present in your body than ever before.", italic=True, color=(36, 22, 19))
    add_body_paragraph(doc2, "Two... moving closer to the root.", italic=True, color=(36, 22, 19))
    add_body_paragraph(doc2, "One... you are there. Safe, focused, receptive.", italic=True, color=(36, 22, 19))
    
    add_custom_heading(doc2, "Section 3: Somatic Regression (8:00 - 14:00)", 2)
    add_highlight_box(doc2, "Speaking Notes", "Deeply emotional, gentle, warm container.")
    
    add_body_paragraph(doc2, "Now, recall that physical feeling of pressure, tension, or defensiveness you logged in your journal. Feel it in your body now. Where is it? In your chest? In your throat?", italic=True, color=(36, 22, 19))
    add_body_paragraph(doc2, "Let that feeling act as a thread. We are following it back, back, back, through time. Back past last year... past college... past high school... down to a childhood address.", italic=True, color=(36, 22, 19))
    add_body_paragraph(doc2, "Locate the child who first felt this way. Look at them. See what they are wearing. See the room.", italic=True, color=(36, 22, 19))
    add_body_paragraph(doc2, "What are they afraid of? What script did they write to keep themselves safe in that room? Was it to hide? Was it to perform? Was it to build an empire so no one could touch them?", italic=True, color=(36, 22, 19))
    add_body_paragraph(doc2, "Step into the room as your adult self. Look at this child. Place your hand gently on their shoulder, and tell them what they have needed to hear for twenty years:", italic=True, color=(36, 22, 19))
    add_body_paragraph(doc2, "'I am here now. You don't have to carry this anymore. The danger is over. You did a beautiful job keeping us safe, but I am the adult now, and I've got this. You are allowed to play. You are allowed to rest.'", italic=True, color=(107, 39, 55))
    add_body_paragraph(doc2, "Feel that child step into your chest, releasing the grip. Feel the physical space opening up in your lungs.", italic=True, color=(36, 22, 19))
    
    add_custom_heading(doc2, "Section 4: Somatic Integration (14:00 - 18:00)", 2)
    add_highlight_box(doc2, "Speaking Notes", "Rises in volume slightly. Confident, grounding.")
    
    add_body_paragraph(doc2, "With every breath, wire this safety in. Your nervous system is updating. The old threat response is dissolving. You no longer need to perform to be safe. You are safe because you exist.", italic=True, color=(36, 22, 19))
    add_body_paragraph(doc2, "In a moment, I will count from one to five, bringing you back to the room.", italic=True, color=(36, 22, 19))
    add_body_paragraph(doc2, "One... feeling the weight of your feet on the floor.", italic=True, color=(36, 22, 19))
    add_body_paragraph(doc2, "Two... carrying this deep somatic peace with you.", italic=True, color=(36, 22, 19))
    add_body_paragraph(doc2, "Three... drawing in a fresh, energizing breath.", italic=True, color=(36, 22, 19))
    add_body_paragraph(doc2, "Four... moving your fingers and toes.", italic=True, color=(36, 22, 19))
    add_body_paragraph(doc2, "Five... eyes open, fully back, fully integrated.", italic=True, color=(36, 22, 19))
    
    add_custom_heading(doc2, "Section 5: The Testimonial Anchor & Funnel Hook (18:00 - 20:00)", 2)
    add_highlight_box(doc2, "Speaking Notes", "Conversational, warm, direct tone.")
    
    add_body_paragraph(doc2, "Welcome back. Take a moment to feel your body. Feel the silence in your mind.", italic=True, color=(36, 22, 19))
    add_body_paragraph(doc2, "What you are experiencing right now is not just relaxation—it is a physical update to your somatic patterns. While this shift is raw, while it is real, let us anchor it.", italic=True, color=(36, 22, 19))
    add_body_paragraph(doc2, "Open your journal to Page 7. Scan the portal code. Speak your truth. Record a brief, 2-minute video for me. Share how you felt coming in, what occurred during the regression, and how your chest and body feel right now.", italic=True, color=(36, 22, 19))
    add_body_paragraph(doc2, "Sharing your shift cements this state. It tells your subconscious that the safety you found is real. Once you submit your reflection, the portal will instantly unlock your R200 LoveBetter Assessment Voucher so we can map your full attachment loops, along with your somatic reset audio.", italic=True, color=(36, 22, 19))
    add_body_paragraph(doc2, "I look forward to hearing your voice. Speak soon.", italic=True, color=(36, 22, 19))
    
    # Save doc 2
    doc2_path = os.path.join(output_dir, "FOLA_Inner_Child_Regression_Hypnosis_Script.docx")
    doc2.save(doc2_path)
    print(f"✅ Hypnosis Script DOCX Successfully Saved to: {doc2_path}")
    print("🎉 All DOCX files successfully generated.")

if __name__ == "__main__":
    main()
